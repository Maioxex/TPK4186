import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

class report():

    def __init__(self):
        self.document = Document()
    
    def getdocument(self):
        return self.document
    
    def addTitle(self, title):
        self.document.add_heading(title, level=0) 
    
    def addPlot(self, plot):
        self.document.add_picture(plot, width=Inches(6))
    
    def save(self, filename):
        self.document.save(filename)
    
    def addSection(self, title):
        self.document.add_heading(title, level=1)
        
    def addParagraph(self, text):
        self.document.add_paragraph(text)
       
    def addPicture(self, image):
        self.document.add_picture(image, width=Inches(6))
    
    #this is the shit for task 7
    def createtablestatdoc(self, whiteresults, blackresults, total):
        lists = [total, whiteresults, blackresults]
        self.document.add_table(rows=4, cols=4)
        for i in range (1,4):
            for j in range (1,4):
                self.document.tables[0].cell(i, j).text = f"{lists[i-1][j-1]}"
        self.document.tables[0].cell(1, 0).text = "Total"
        self.document.tables[0].cell(0, 1).text = "Wins"
        self.document.tables[0].cell(0, 2).text = "Losses"
        self.document.tables[0].cell(0, 3).text = "Draws"
        self.document.tables[0].cell(2, 0).text = "White"
        self.document.tables[0].cell(3, 0).text = "Black"

    def createtabletma4240doc(self, avarage,std,whiteavg,blackavg,whitestd,blackstd, winstd, winavg, losstd, losavg):
        self.document.add_table(rows=6, cols=3)
        self.document.tables[1].cell(0, 1).text = "Average"
        self.document.tables[1].cell(0, 2).text = "Standard Deviation"
        self.document.tables[1].cell(1, 1).text = str(avarage)
        self.document.tables[1].cell(1, 2).text = str(std)
        self.document.tables[1].cell(1, 0).text = "Total"
        self.document.tables[1].cell(2, 0).text = "White"
        self.document.tables[1].cell(3, 0).text = "Black"
        self.document.tables[1].cell(2, 1).text = str(whiteavg)
        self.document.tables[1].cell(2, 2).text = str(whitestd)
        self.document.tables[1].cell(3, 1).text = str(blackavg)
        self.document.tables[1].cell(3, 2).text = str(blackstd)
        self.document.tables[1].cell(4, 0).text = "Wins"
        self.document.tables[1].cell(5, 0).text = "Losses"
        self.document.tables[1].cell(4, 1).text = str(winavg)
        self.document.tables[1].cell(4, 2).text = str(winstd)
        self.document.tables[1].cell(5, 1).text = str(losavg)
        self.document.tables[1].cell(5, 2).text = str(losstd)





# # Generate some random data and plot it
# x = [1, 2, 3, 4, 5]
# y = [10, 8, 6, 4, 2]
# plt.plot(x, y)
# plt.savefig('my_plot.png')
# # Create a new Word document
# document = report()
# document.addTitle("My Report")
# document.addSection("Section 1")
# document.addParagraph("This is a paragraph")
# document.addPicture("my_plot.png")
# document.save("my_report.docx")